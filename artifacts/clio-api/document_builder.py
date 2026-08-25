"""
document_builder.py
====================
Document Generation Engine for DOCX and PDF work products with diagonal watermark.
"""

from __future__ import annotations

import io
import logging
from typing import Any, Literal

from config import DOCUMENT_FOOTER, ETHICS_DISCLAIMERS

logger = logging.getLogger(__name__)

def build_document(
    note_dict: dict[str, Any],
    fmt: Literal["docx", "pdf"] = "docx",
    firm_name: str = "Example Law Firm, PLLC",
    matter_label: str = "",
) -> tuple[str, bytes, str]:
    if fmt == "docx":
        return _build_docx(note_dict, firm_name, matter_label)
    else:
        return _build_pdf(note_dict, firm_name, matter_label)


def _build_docx(note_dict: dict[str, Any], firm_name: str, matter_label: str) -> tuple[str, bytes, str]:
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        # Fallback text bytes if docx not installed
        txt = f"{firm_name}\nMATTER: {matter_label}\nSUBJECT: {note_dict.get('suggested_subject')}\n\nSUMMARY:\n{note_dict.get('summary')}"
        return "work_product.txt", txt.encode("utf-8"), "text/plain"

    doc = docx.Document()
    doc.add_heading(firm_name, level=1)
    if matter_label:
        doc.add_paragraph(f"Matter: {matter_label}")
    doc.add_paragraph(f"Subject: {note_dict.get('suggested_subject')}")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(note_dict.get("summary", ""))

    doc.add_heading("Key Facts", level=2)
    for f in note_dict.get("facts", []):
        doc.add_paragraph(f, style="List Bullet")

    doc.add_paragraph(DOCUMENT_FOOTER)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return f"legal_note_{matter_label or 'draft'}.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_pdf(note_dict: dict[str, Any], firm_name: str, matter_label: str) -> tuple[str, bytes, str]:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        txt = f"{firm_name}\nPDF FALLBACK\n{note_dict.get('summary')}"
        return "work_product.txt", txt.encode("utf-8"), "text/plain"

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)

    # Background Watermark
    c.saveState()
    c.setFont("Helvetica-Bold", 36)
    c.setFillColorRGB(0.9, 0.2, 0.2, 0.15)
    c.rotate(45)
    c.drawString(150, 100, "AI DRAFT - ATTORNEY REVIEW REQUIRED")
    c.restoreState()

    # Content
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 750, firm_name)
    c.setFont("Helvetica", 12)
    c.drawString(50, 730, f"Subject: {note_dict.get('suggested_subject')}")
    c.drawString(50, 710, f"Matter: {matter_label or 'General'}")

    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, 680, "Summary:")
    c.setFont("Helvetica", 10)

    text_obj = c.beginText(50, 660)
    summary_text = note_dict.get("summary", "")
    for line in summary_text.split("\n"):
        text_obj.textLine(line[:90])
    c.drawText(text_obj)

    c.setFont("Helvetica-Oblique", 8)
    c.drawString(50, 40, ETHICS_DISCLAIMERS["primary"])
    c.showPage()
    c.save()

    buf.seek(0)
    return f"legal_note_{matter_label or 'draft'}.pdf", buf.getvalue(), "application/pdf"
